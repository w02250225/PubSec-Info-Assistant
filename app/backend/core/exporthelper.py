import logging
import io
import re
import time
import docx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from azure.storage.blob.aio import ContainerClient
from openai import AsyncOpenAI
from quart import request, session
from typing import List, Dict, Union

async def export_to_blob(container_client: ContainerClient,
                   openai_client: AsyncOpenAI,
                   gpt_deployment: str,
                   gpt_model_name: str
                   ):
    try:
        request_data = await request.json
        user_id = session.get('user_data', {}).get('userPrincipalName') or "Unknown User"

        title = await generate_document_title(request_data["question"],
                                              openai_client,
                                              gpt_deployment,
                                              gpt_model_name)

        export = await create_docx(title,
                                   request_data["answer"],
                                   request_data["citations"])

        await upload_to_blob(export,
                             request_data["request_id"],
                             user_id,
                             container_client)

        # Use title to generate file name
        # Make sure it's safe for Windows
        file_name = f"{sanitize_filename(title)}.docx"

        # Go back to the start of the stream
        export.seek(0)

        return file_name, export

    except Exception as ex:
        logging.exception("Exception in export_to_blob")
        return str(ex)


async def generate_document_title(question: str,
                                  openai_client: AsyncOpenAI,
                                  gpt_deployment: str,
                                  gpt_model_name: str
                                  ) -> str:

    messages = [
        {"role": "system", "content": "You are a helpful AI that generates document titles."},
        {"role": "user", "content": f"Generate a short (less than 10 words) document title for the following question or request: \"{question}\""},
    ]

    chat_completion = await openai_client.chat.completions.create(
                            messages = messages,
                            # Azure Open AI takes the deployment name as the model name
                            model = gpt_deployment if gpt_deployment else gpt_model_name,
                            temperature = 0.4,
                            max_tokens = 20,
                            n = 1)
    
    if chat_completion.choices:
        first_choice = chat_completion.choices[0]
        title = first_choice.message.content.replace("\"", "")
        return title
    else:
        return "No title generated"


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rpr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rpr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    run = paragraph.add_run()
    run._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    # pylint: disable=E1101  # Disable "no-member" error for 'HYPERLINK' in MSO_THEME_COLOR_INDEX
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True

    return hyperlink


def create_new_doc(title):
    doc = Document()

    # Add header
    paragraph = doc.sections[0].header.paragraphs[0]
    paragraph.text = "Generated with Coeus - SENSITIVE - DRAFT ONLY"

    # Add title
    if len(doc.paragraphs) == 0:
        doc.add_paragraph()
    para = doc.paragraphs[0]
    para.add_run(title).font.size = Pt(22)
    return doc


def add_html_to_docx(html: str, doc, heading=None):

    html = html.split('<')
    html = [html[0]] + ['<'+l for l in html[1:]]
    tags = []

    para = doc.add_paragraph()
    if heading:
        para.add_run(heading).font.size = Pt(14)
        para = doc.add_paragraph()

    for run in html:
        tag_change = re.match('(?:<)(.*?)(?:>)', run)
        if tag_change is not None:
            tag_strip = tag_change.group(0)
            tag_change = tag_change.group(1)
            if tag_change.startswith('/'):
                if tag_change.startswith('/a'):
                    tag_change = next(
                        tag for tag in tags if tag.startswith('a '))
                tag_change = tag_change.strip('/')
                tags.remove(tag_change)
            else:
                tags.append(tag_change)
        else:
            tag_strip = ''
        hyperlink = [tag for tag in tags if tag.startswith('a ')]
        if run.startswith('<'):
            run = run.replace(tag_strip, '')
            if hyperlink:
                hyperlink = hyperlink[0]
                hyperlink = re.match(
                    '.*?(?:href=")(.*?)(?:").*?', hyperlink).group(1)
                add_hyperlink(para, run, hyperlink)
            else:
                runner = para.add_run(run)
                if 'b' in tags:
                    runner.bold = True
                if 'u' in tags:
                    runner.underline = True
                if 'i' in tags:
                    runner.italic = True
                if 'h1' in tags:
                    runner.font.size = Pt(18)
                if 'h2' in tags:
                    runner.font.size = Pt(14)
                if 'h3' in tags:
                    runner.font.size = Pt(12)
                if 'sup' in tags:
                    runner.font.superscript = True
        else:
            para.add_run(run)


async def create_docx(title: str, answer: str, citations: List[Dict[str, Union[str, int]]]):
    try:
        soup_answer = BeautifulSoup(answer, "lxml")

        # Remove citation anchor in answer text
        for a in soup_answer.findAll('a'):
            a.replaceWithChildren()

        doc = create_new_doc(title)

        # Add answer to the document
        # Add spaces between the citations
        html_answer = str(soup_answer.body).replace(
            '</sup><sup>', '</sup> <sup>')
        add_html_to_docx(html_answer, doc)

        # Construct citations from the citations array
        server_hostname = request.url_root
        html_citations = ''
        for citation in citations:
            link = f'{server_hostname}/#/ViewDocument?documentName={citation["sourceFile"]}&pageNumber={citation["pageNumber"]}'
            html_citations += f'<a href="{link}">{citation["label"]}</a>\n'

        add_html_to_docx(html_citations, doc, "Citations")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        return stream

    except Exception as ex:
        logging.exception("Exception in create_docx")
        return str(ex)


async def upload_to_blob(input_stream: io.BytesIO,
                         request_id: str,
                         user_id: str,
                         container_client: ContainerClient) -> str:
    try:
        timestamp = time.strftime("%Y%m%d%H%M%S")
        blob_name = f"{user_id}/{timestamp}_{request_id}.docx"
        await container_client.get_blob_client(blob_name).upload_blob(
            input_stream, 
            blob_type="BlockBlob"
            )

        return blob_name

    except Exception as ex:
        logging.exception("Exception in upload_to_blob")
        return str(ex)

def sanitize_filename(title: str) -> str:
    # Define a regular expression pattern to match characters not allowed in file names
    invalid_chars = r'[<>:"/\\|?*]'

    # Replace invalid characters with underscores
    sanitized_title = re.sub(invalid_chars, '', title)

    # Remove any leading or trailing spaces
    sanitized_title = sanitized_title.strip()

    # Ensure the filename does not exceed the Windows maximum length (255 characters)
    if len(sanitized_title) > 255:
        sanitized_title = sanitized_title[:255]

    return sanitized_title
